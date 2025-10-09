from fastapi import FastAPI, HTTPException, Query, File, UploadFile, APIRouter, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import pymysql
import os
from dotenv import load_dotenv
from pinecone import Pinecone
from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
import json
import requests
from datetime import datetime
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import RGBColor
import tempfile
from db.chunking.regex_chunks_sop import process_single_pdf, process_pdfs_parallel, SemanticPDFProcessor
from db.pinecone.sop_vectordb import filter_sop_chunks, prepare_sop_documents, embed_documents_in_memory, EMBEDDING_MODEL, BATCH_SIZE
from db.pinecone.to_pinecone import upsert_embeddings_to_pinecone
from auto_education import GMPTrainingService
import time
import pickle


load_dotenv()

app = FastAPI(
    title="GMP-SOP API",
    description="SOP 문서와 GMP 요약을 연결하는 API 서버",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000", "http://localhost:3001"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 환경변수 및 상수
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
INDEX_NAME = "gmp-sop-vectordb"
VECTOR_DIM = 1536
ALLOWED_NAMESPACES = ["sop", "gmp-1st", "gmp-2nd", "old-gmp-1st", "old-gmp-2nd"]

# DB 접속 정보 (Cloudtype 환경)
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'mariadb'),
    'port': int(os.getenv('DB_PORT', 3306)),
    'user': os.getenv('DB_USER'),
    'password': os.getenv('DB_PASSWORD'),
    'database': os.getenv('DB_NAME'),
    'charset': 'utf8mb4'
}

def get_db_connection():
    return pymysql.connect(**DB_CONFIG)

# ---- Pydantic 모델 정의 ----
class AnalysisSummaryModel(BaseModel):
    총_gmp_변경점: int = Field(..., alias="total_gmp_changes")
    영향받는_sop_섹션: int = Field(..., alias="affected_sop_sections")
    분석완료시각: str = Field(..., alias="analyzed_at")

class SopInfoModel(BaseModel):
    sop_id: str
    sop_title: Optional[str] = None
    sop_content: Optional[str] = None
    match_score: Optional[float] = 0

class GmpChangeInfoModel(BaseModel):
    change_id: str
    topic: Optional[str] = None
    old_gmp_content: Optional[str] = ""
    new_gmp_content: Optional[str] = ""
    similarity_score: Optional[float] = 0

class DetailedAnalysisModel(BaseModel):
    sop_info: SopInfoModel
    gmp_change_info: Optional[GmpChangeInfoModel] = None
    change_rationale: Optional[Dict[str, Any]] = {}
    update_recommendation: Optional[str] = ""

class SaveAllRequestModel(BaseModel):
    summary: AnalysisSummaryModel
    detailed_analysis: List[DetailedAnalysisModel]

class SopUpdateModel(BaseModel):
    sop_title: Optional[str] = None
    sop_content: Optional[str] = None
    # 필요하다면 modified, updated_at 등도 추가 가능

class TrainingRequest(BaseModel):
    sop_content: str
    guideline_content: str
    target_audience: str = "GMP 실무자"
    num_questions: int = 15


# ---- DB 저장 함수 ----
def insert_analysis_summary(summary: AnalysisSummaryModel):
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS ANALYSIS_SUMMARY (
                    total_gmp_changes INT,
                    affected_sop_sections INT,
                    created_at VARCHAR(50)
                )
            ''')
            sql = """
                INSERT INTO ANALYSIS_SUMMARY (
                    total_gmp_changes,
                    affected_sop_sections,
                    created_at
                ) VALUES (%s, %s, %s)
            """
            cursor.execute(sql, (
                summary.총_gmp_변경점,
                summary.영향받는_sop_섹션,
                summary.분석완료시각
            ))
        conn.commit()
    finally:
        conn.close()

def insert_sop_data(sop_list: List[DetailedAnalysisModel]):
    create_sql = '''
        CREATE TABLE IF NOT EXISTS SOP (
            sop_id VARCHAR(100) PRIMARY KEY,
            sop_title VARCHAR(255),
            sop_content TEXT,
            is_current BOOLEAN DEFAULT FALSE,
            created_at DATETIME,
            updated_at DATETIME
        )
    '''
    insert_sql = """
        INSERT INTO SOP (
            sop_id,
            sop_title,
            sop_content,
            created_at,
            updated_at
        ) VALUES (%s, %s, %s, NOW(), NOW())
        ON DUPLICATE KEY UPDATE
            sop_title = VALUES(sop_title),
            sop_content = VALUES(sop_content),
            updated_at = NOW()
    """
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute(create_sql)
            for item in sop_list:
                sop_info = item.sop_info
                if not sop_info.sop_id or not sop_info.sop_content:
                    continue
                cursor.execute(insert_sql, (
                    sop_info.sop_id,
                    sop_info.sop_title,
                    sop_info.sop_content
                ))
            conn.commit()
    finally:
        conn.close()

def insert_gmp_data(gmp_list):
    create_sql = '''
        CREATE TABLE IF NOT EXISTS GMP (
            gmp_id VARCHAR(50) PRIMARY KEY,
            topic VARCHAR(255),             
            gmp_content TEXT,              
            similarity_score FLOAT,         
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    '''
    insert_sql = """
        INSERT INTO GMP (
            gmp_id,
            topic,
            gmp_content,
            similarity_score,
            created_at
        ) VALUES (%s, %s, %s, %s, NOW())
        ON DUPLICATE KEY UPDATE
            topic = VALUES(topic),
            gmp_content = VALUES(gmp_content),
            similarity_score = VALUES(similarity_score)
    """
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute(create_sql)
            for gmp_item in gmp_list:
                gmp_changes = gmp_item.gmp_changes or []
                top_gmp_changes = sorted(gmp_changes, key=lambda x: x.get('match_score', 0), reverse=True)[:5]

                for gmp_info in top_gmp_changes:
                    gmp_id = gmp_info.get('change_id')
                    topic = gmp_info.get('topic')
                    old_content = gmp_info.get('old_gmp_summary', '')
                    new_content = gmp_info.get('new_gmp_summary', '')
                    gmp_content = f"OLD:\n{old_content}\n\nNEW:\n{new_content}"
                    similarity_score = gmp_info.get('similarity_score', 0)
                    if gmp_id is None or not gmp_content:
                        continue
                    cursor.execute(insert_sql, (
                        gmp_id,
                        topic,
                        gmp_content,
                        similarity_score
                    ))
            conn.commit()
    finally:
        conn.close()

def insert_sop_gmp_link(sop_list):
    create_sql = '''
        CREATE TABLE IF NOT EXISTS SOP_GMP_LINK (
            link_id BIGINT AUTO_INCREMENT PRIMARY KEY,
            sop_id VARCHAR(50) NOT NULL,
            gmp_id VARCHAR(50) NOT NULL,
            match_score FLOAT,              
            change_rationale JSON,          
            update_recommendation TEXT,  
            completed VARCHAR(20) NOT NULL DEFAULT '미처리', 
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (sop_id) REFERENCES SOP(sop_id) ON DELETE CASCADE,
            FOREIGN KEY (gmp_id) REFERENCES GMP(gmp_id) ON DELETE CASCADE,
            UNIQUE KEY unique_sop_gmp (sop_id, gmp_id)
        )
    '''
    insert_sql = """
        INSERT INTO SOP_GMP_LINK (
            sop_id,
            gmp_id,
            match_score,
            change_rationale,
            update_recommendation,
            completed,
            created_at
        ) VALUES (%s, %s, %s, %s, %s, %s, NOW())
        ON DUPLICATE KEY UPDATE
            match_score = VALUES(match_score),
            change_rationale = VALUES(change_rationale),
            update_recommendation = VALUES(update_recommendation),
            completed = VALUES(completed)
    """
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute(create_sql)
            for sop_item in sop_list:
                sop_info = sop_item.sop_info
                gmp_changes = sop_item.gmp_changes or []
                rationale = sop_item.change_rationale or {}
                update_recommendation = sop_item.update_recommendation or ''
                completed_status = '미처리'
                sop_id = sop_info.sop_id if sop_info else None

                top_gmp_changes = sorted(gmp_changes, key=lambda x: x.get('match_score', 0), reverse=True)[:5]

                for gmp_info in top_gmp_changes:
                    gmp_id = gmp_info.get('change_id')
                    match_score = gmp_info.get('match_score', 0)
                    if sop_id is None or gmp_id is None:
                        continue
                    change_rationale_json = json.dumps(rationale, ensure_ascii=False)
                    cursor.execute(insert_sql, (
                        sop_id,
                        gmp_id,
                        match_score,
                        change_rationale_json,
                        update_recommendation,
                        completed_status
                    ))
            conn.commit()
    finally:
        conn.close()

def drop_all_tables(DB_CONFIG):
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute("SET FOREIGN_KEY_CHECKS = 0;")
            cursor.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = %s;", (DB_CONFIG['database'],))
            tables = cursor.fetchall()
            for (table_name,) in tables:
                print(f"Dropping table: {table_name}")
                cursor.execute(f"DROP TABLE IF EXISTS `{table_name}`;")
            cursor.execute("SET FOREIGN_KEY_CHECKS = 1;")
        conn.commit()
        print("모든 테이블 삭제 완료")
    finally:
        conn.close()

def create_all_tables(DB_CONFIG):
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor() as cursor:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ANALYSIS_SUMMARY (
                    total_gmp_changes INT,
                    affected_sop_sections INT,
                    created_at VARCHAR(50)
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS SOP (
                    sop_id VARCHAR(100) PRIMARY KEY,
                    sop_title VARCHAR(255),
                    sop_content TEXT,
                    is_current BOOLEAN DEFAULT FALSE,
                    created_at DATETIME,
                    updated_at DATETIME
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS GMP (
                    gmp_id VARCHAR(50) PRIMARY KEY,
                    topic VARCHAR(255),
                    gmp_content TEXT,
                    similarity_score FLOAT,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS SOP_GMP_LINK (
                    link_id BIGINT AUTO_INCREMENT PRIMARY KEY,
                    sop_id VARCHAR(50) NOT NULL,
                    gmp_id VARCHAR(50) NOT NULL,
                    match_score FLOAT,
                    change_rationale JSON,
                    update_recommendation TEXT,
                    completed VARCHAR(20) NOT NULL DEFAULT '미처리',
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (sop_id) REFERENCES SOP(sop_id) ON DELETE CASCADE,
                    FOREIGN KEY (gmp_id) REFERENCES GMP(gmp_id) ON DELETE CASCADE,
                    UNIQUE KEY unique_sop_gmp (sop_id, gmp_id)
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS SOP_MODIFIED (
                    modified_content_id BIGINT AUTO_INCREMENT PRIMARY KEY,
                    sop_id VARCHAR(50) NOT NULL,
                    sop_title VARCHAR(255) NOT NULL,
                    sop_content TEXT NOT NULL,
                    educated BOOLEAN DEFAULT FALSE,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (sop_id) REFERENCES SOP(sop_id) ON DELETE CASCADE
                )
            """)
        conn.commit()
        print("모든 테이블 생성 완료")
    finally:
        conn.close()

# ---- FastAPI 엔드포인트 ----
# ---- 파일 업로드 엔드포인트 ----
@app.post("/upload_json")
async def upload_json(file: UploadFile = File(...)):
    try:
        drop_all_tables(DB_CONFIG)
        create_all_tables(DB_CONFIG)
        contents = await file.read()
        data = json.loads(contents.decode('utf-8'))
        summary = data.get('summary', {})
        summary['analyzed_at'] = datetime.utcnow().isoformat()
        # 한글 키가 있으면 영문 키로 복사
        if '총_gmp_변경점' in summary:
            summary['total_gmp_changes'] = summary['총_gmp_변경점']
        if '영향받는_sop_섹션' in summary:
            summary['affected_sop_sections'] = summary['영향받는_sop_섹션']

        # 필수 키가 없으면 에러 메시지 반환
        required_keys = ['total_gmp_changes', 'affected_sop_sections', 'analyzed_at']
        for key in required_keys:
            if key not in summary:
                raise HTTPException(status_code=400, detail=f"'{key}' 필드가 summary에 없습니다.")

        summary_model = AnalysisSummaryModel(**summary)
        detailed = [DetailedAnalysisModel(**item) for item in data['detailed_analysis']]
        insert_analysis_summary(summary_model)
        insert_sop_data(detailed)
        insert_gmp_data(detailed)
        insert_sop_gmp_link(detailed)
        return {"result": "파일 업로드 및 데이터 저장 성공!"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

pinecone_index = None

@app.on_event("startup")
def startup_event():
    global pinecone_index
    try:
        if not PINECONE_API_KEY:
            raise ValueError("PINECONE_API_KEY 환경변수가 필요합니다.")
        pc = Pinecone(api_key=PINECONE_API_KEY)
        pinecone_index = pc.Index(INDEX_NAME)
    except Exception as e:
        print(f"Pinecone 초기화 실패: {e}")
        pinecone_index = None

def test_connection():
    try:
        conn = get_db_connection()
        with conn.cursor() as cursor:
            cursor.execute("SELECT 1")
            result = cursor.fetchone()
        conn.close()
        if result:
            return {"status": "연결 성공"}
        else:
            raise Exception("쿼리 결과 없음")
    except Exception as e:
        return {"status": "연결 실패", "error": str(e)}

@app.get("/")
def read_root():
    return {
        "message": "GMP-SOP API 서버 정상 실행 중",
        "version": "1.0.0",
        "status": "running"
    }

@app.get("/get_sop_text")
def get_sop_text(
    namespace: str = Query("sop"),
    top_k: int = Query(10, ge=1, le=100),
    offset: int = Query(0, ge=0),
):
    if pinecone_index is None:
        raise HTTPException(status_code=500, detail="Pinecone 인덱스가 초기화되어 있지 않습니다.")
    if namespace not in ALLOWED_NAMESPACES:
        raise HTTPException(status_code=400, detail=f"지원되는 네임스페이스는 {ALLOWED_NAMESPACES} 입니다.")
    try:
        dummy_vector = [0.0] * VECTOR_DIM
        results = pinecone_index.query(
            vector=dummy_vector,
            top_k=top_k + offset,
            namespace=namespace,
            include_metadata=True
        )
        matches = results.get("matches", [])[offset:]
        data = [match.get("metadata", {}) for match in matches]
        return {
            "namespace": namespace,
            "count": len(data),
            "data": data,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Pinecone 쿼리 오류: {e}")

@app.get("/get_sop_to_update")
def get_sop_to_update():
    conn = get_db_connection()
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            sql = """
                SELECT sop_id, sop_title, sop_content, created_at, updated_at
                FROM SOP
                ORDER BY updated_at DESC
            """
            cursor.execute(sql)
            sop_data = cursor.fetchall()
        return {"sop_data": sop_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB 조회 오류: {e}")
    finally:
        conn.close()

@app.get("/get_changed_gmp")
def get_changed_gmp():
    conn = get_db_connection()
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            sql = """
                SELECT gmp_id, topic, gmp_content, similarity_score, created_at
                FROM GMP
                ORDER BY created_at DESC
            """
            cursor.execute(sql)
            gmp_data = cursor.fetchall()
        return {"gmp_data": gmp_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB 조회 오류: {e}")
    finally:
        conn.close()

@app.get("/get_sop_gmp_link")
def get_sop_gmp_link():
    conn = get_db_connection()
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            sql = """
                SELECT sop_id, gmp_id, match_score, change_rationale, update_recommendation, completed, created_at
                FROM SOP_GMP_LINK
                ORDER BY created_at DESC
            """
            cursor.execute(sql)
            link_data = cursor.fetchall()
        return {"link_data": link_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB 조회 오류: {e}")
    finally:
        conn.close()

@app.patch("/sop/{sop_id}")
def update_sop(sop_id: str, update: SopUpdateModel):
    update_data = update.dict(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="수정할 내용이 없습니다.")
    
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    update_data['updated_at'] = now_str
    update_data['modified'] = True  # SOP 테이블의 modified 열 True로 업데이터 -> 수정 이력 관리

    set_clause = ", ".join([f"{k}=%s" for k in update_data.keys()])
    values = list(update_data.values())
    values.append(sop_id)

    conn = get_db_connection()
    try:
        with conn.cursor() as cursor:
            # 1. SOP 테이블 업데이트
            sql_update_sop = f"UPDATE SOP SET {set_clause} WHERE sop_id=%s"
            cursor.execute(sql_update_sop, values)

            # 2. SOP_MODIFIED 테이블에 수정 내용 insert
            # 수정할 제목과 내용은 update_data에 있을 수 있으니 예외 처리

            # 테이블 생성 쿼리
            create_modified_sql = '''
                CREATE TABLE IF NOT EXISTS SOP_MODIFIED (
                    modified_content_id BIGINT AUTO_INCREMENT PRIMARY KEY,
                    sop_id VARCHAR(50) NOT NULL,
                    sop_title VARCHAR(255) NOT NULL,
                    sop_content TEXT NOT NULL,
                    educated BOOLEAN DEFAULT FALSE,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                    FOREIGN KEY (sop_id) REFERENCES SOP(sop_id) ON DELETE CASCADE
                )
            '''
            title = update_data.get('sop_title')
            content = update_data.get('sop_content')

            if title is not None and content is not None:
                cursor.execute(create_modified_sql)
                sql_insert_modified = """
                    INSERT INTO SOP_MODIFIED (sop_id, sop_title, sop_content, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s)
                """
                cursor.execute(sql_insert_modified, (sop_id, title, content, now_str, now_str))

        conn.commit()
    finally:
        conn.close()

    return {"result": "SOP 수정 및 수정내용 저장 성공"}


# 수정한 SOP 문서로 출력
@app.get("/export_sop_docx")
def export_sop_docx():
    # 1. Pinecone에서 SOP 전문 메타데이터 모두 조회
    namespace = "sop"
    top_k = 10000  # 충분히 큰 값으로 전체 조항 조회
    dummy_vector = [0.0] * VECTOR_DIM
    results = pinecone_index.query(
        vector=dummy_vector,
        top_k=top_k,
        namespace=namespace,
        include_metadata=True
    )
    matches = results.get("matches", [])
    data = [match.get("metadata", {}) for match in matches]

    # 2. title별로 그룹화, 각 그룹 내에서 chunk_index로 정렬
    title_groups = defaultdict(list)
    for item in data:
        title = item.get("title", "(제목 없음)")
        title_groups[title].append(item)
    # 각 title 그룹 내에서 chunk_index기준으로 정렬
    for title in title_groups:
        title_groups[title].sort(key=lambda x: int(x.get("chunk_index", 0)))
    
    # 3. DB에서 SOP_MODIFIED 테이블의 최신 수정안 조회 (sop_id 기준)
    conn = get_db_connection()
    sop_mod_map = {}
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            for item in data:
                sop_id = item.get("id")
                if sop_id:
                    cursor.execute("""
                        SELECT sop_title, sop_content
                        FROM SOP_MODIFIED
                        WHERE sop_id=%s
                        ORDER BY updated_at DESC
                        LIMIT 1
                    """, (sop_id,))
                    mod_row = cursor.fetchone()
                    if mod_row:
                        sop_mod_map[sop_id] = mod_row
    finally:
        conn.close()

    # 4. 워드 문서 생성
    doc = Document()
    doc.add_heading("SOP 전문", 0)
    first_title = True
    for title, items in title_groups.items():
        if not first_title:
            doc.add_page_break()  # title별 페이지 구분
        first_title = False
        # title(문서 단위) 헤딩
        heading = doc.add_heading(title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # 진한 파란색 등으로 구분
        heading.runs[0].font.bold = True
        for item in items:
            section_title = item.get("section_title", "")
            text = item.get("text", "")
            sop_id = item.get("id")
            modified = item.get("modified", False)
            # 조항 제목
            para = doc.add_paragraph(section_title)
            para.runs[0].font.bold = True
            # 조항 본문
            doc.add_paragraph(text)
            # 수정안 표시
            if modified and sop_id in sop_mod_map:
                mod = sop_mod_map[sop_id]
                mod_paragraph = doc.add_paragraph()
                mod_run = mod_paragraph.add_run(f"[수정안] {mod['sop_title']}\n{mod['sop_content']}")
                mod_run.font.bold = True  # 굵게
                mod_run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색 등으로 구분
        doc.add_paragraph("")  # 빈 줄로 구분

    # 5. 임시 파일로 저장 및 반환
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    filename = f"SOP_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    response = FileResponse(tmp_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    def cleanup():
        import os
        os.remove(tmp_path)
    response.background = BackgroundTasks()
    response.background.add_task(cleanup)
    return response

def export_gmp_docx(sop_id: str):
    conn = get_db_connection()
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # SOP_GMP_LINK에서 sop_id로 연결된 gmp_id, 근거 등 조회
            cursor.execute("""
                SELECT gmp_id, change_rationale, update_recommendation
                FROM SOP_GMP_LINK
                WHERE sop_id=%s AND completed='처리'
            """, (sop_id,))
            link_data = cursor.fetchall()
            gmp_ids = [row['gmp_id'] for row in link_data]
            # GMP 테이블에서 전문 조회
            if gmp_ids:
                format_strings = ','.join(['%s'] * len(gmp_ids))
                cursor.execute(f"""
                    SELECT gmp_id, topic, gmp_content
                    FROM GMP
                    WHERE gmp_id IN ({format_strings})
                """, tuple(gmp_ids))
                gmp_data = {row['gmp_id']: row for row in cursor.fetchall()}
            else:
                gmp_data = {}
    finally:
        conn.close()

    # 워드 문서 생성
    doc = Document()
    doc.add_heading(f"SOP({sop_id})와 연결된 GMP 전문", 0)
    for link in link_data:
        gmp = gmp_data.get(link['gmp_id'])
        doc.add_heading(gmp['topic'] if gmp else link['gmp_id'], level=1)
        doc.add_paragraph(gmp['gmp_content'] if gmp else "(GMP 내용 없음)")
        doc.add_paragraph(f"근거: {link['change_rationale']}")
        doc.add_paragraph(f"업데이트 권장사항: {link['update_recommendation']}")
        doc.add_paragraph("")

    # 임시 파일로 저장 및 반환
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    filename = f"GMP_{sop_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    response = FileResponse(tmp_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    from fastapi import BackgroundTasks
    def cleanup():
        os.remove(tmp_path)
    response.background = BackgroundTasks()
    response.background.add_task(cleanup)
    return response

router = APIRouter()

# 수정한 SOP 업로드 -> 청킹 및 임베딩하여 pinecone에 저장
@router.post("/upload_sop_modified")
async def upload_sop_modified(file: UploadFile = File(...)):
    # 파일 확장자 체크
    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        raise HTTPException(status_code=400, detail="지원하는 파일 형식은 PDF 또는 Word입니다.")
    # 파일 저장 (임시 경로)
    temp_path = f"/tmp/{file.filename}"
    with open(temp_path, "wb") as f:
        f.write(await file.read())
    # PDF 청킹 (텍스트 추출 + 청킹)
    chunks = process_single_pdf(temp_path)  # 바로 청킹 결과 리스트 반환
    if not chunks:
        raise HTTPException(status_code=400, detail="청킹 결과 없음")
    filtered = filter_sop_chunks(chunks, jurisdiction_filter=None)
    documents = prepare_sop_documents(filtered)
    # 임베딩 결과를 임시 변수에 저장
    embedding_results = embed_documents_in_memory(documents, EMBEDDING_MODEL, BATCH_SIZE)
    # 이후 embedding_results를 바로 Pinecone 등으로 넘기면 됨
    metadata_list = [doc.metadata for doc in documents]  # 메타데이터 리스트 추출
    upsert_embeddings_to_pinecone(
        index_name="gmp-sop-vectordb",  
        embeddings=embedding_results,   # 임베딩 결과 리스트
        metadata=metadata_list,         # 메타데이터 리스트
        namespace="sop-2",              # 네임스페이스 지정
        dimension=1536,                 # 임베딩 차원
        reset=False,                    # 필요시 True로
        batch_size=BATCH_SIZE
    )
    return {"result": "업로드 및 임베딩+Pinecone 저장 성공", "chunks": len(documents)}

app.include_router(router)

edu_service = GMPTrainingService(model="gpt-3.5-turbo")

# 교육 생성 엔드포인트
@router.post("/api/edu/bulk-generate")
def bulk_generate_edu(num_questions: int = 15, target_audience: str = "GMP 실무자"):
    conn = pymysql.connect(**DB_CONFIG)
    try:
        with conn.cursor(pymysql.cursors.DictCursor) as cursor:
            # 1. 교육되지 않은 SOP_MODIFIED 전체 조회
            cursor.execute("""
                SELECT modified_content_id, sop_id, sop_title, sop_content
                FROM SOP_MODIFIED
                WHERE educated = FALSE
            """)
            mod_rows = cursor.fetchall()
            if not mod_rows:
                raise HTTPException(status_code=404, detail="교육 대상 SOP가 없습니다.")
            results = []
            updated_ids = []
            for mod_row in mod_rows:
                sop_id = mod_row["sop_id"]
                mod_title = mod_row["sop_title"]
                mod_content = mod_row["sop_content"]
                modified_content_id = mod_row["modified_content_id"]

                # 원본 SOP 가져오기
                cursor.execute("""
                    SELECT sop_title, sop_content
                    FROM SOP
                    WHERE sop_id=%s
                    LIMIT 1
                """, (sop_id,))
                orig_row = cursor.fetchone()
                if not orig_row:
                    continue
                orig_title = orig_row["sop_title"]
                orig_content = orig_row["sop_content"]

                # GMP 변경 근거 조회 (sop_id와 연동된 gmp_id 목록 통해 조회)
                cursor.execute("""
                    SELECT topic, gmp_content
                    FROM GMP
                    WHERE gmp_id IN (
                        SELECT gmp_id FROM SOP_GMP_LINK WHERE sop_id=%s
                    )
                """, (sop_id,))
                gmp_rows = cursor.fetchall()
                # GMP 근거 텍스트 결합
                gmp_contents = "\n\n".join(f"주제: {g['topic']}\n내용: {g['gmp_content']}" for g in gmp_rows) if gmp_rows else ""

                # SOP와 GMP를 하나로 통합하여 sop_content로 전달
                combined_sop_content = (
                    f"=== 변경 전 SOP ===\n제목: {orig_title}\n{orig_content}\n\n"
                    f"=== 변경 후 SOP ===\n제목: {mod_title}\n{mod_content}\n\n"
                    f"=== 관련 GMP 변경 근거 ===\n{gmp_contents}"
                )

                # 가이드라인 텍스트 (필요시 DB 등에서 가져와서 변경 가능)
                guideline_text = "21 CFR Part 211 등 관련 가이드라인 내용"

                # GMPTrainingService 호출 (프롬프트는 서비스 내부에서 처리)
                result = edu_service.generate_training_package(
                    sop_content=combined_sop_content,
                    guideline_content=guideline_text,
                    target_audience=target_audience,
                    num_questions=num_questions
                )

                results.append({
                    "modified_content_id": modified_content_id,
                    "sop_id": sop_id,
                    "training": result
                })
                updated_ids.append(modified_content_id)

            # 2. SOP_MODIFIED educated=True로 업데이트
            if updated_ids:
                format_strings = ",".join(["%s"] * len(updated_ids))
                cursor.execute(f"""
                    UPDATE SOP_MODIFIED SET educated=TRUE WHERE modified_content_id IN ({format_strings})
                """, tuple(updated_ids))
            conn.commit()
        return {"result": "교육자료 생성 및 SOP_MODIFIED 업데이트 성공", "data": results}
    finally:
        conn.close()

if __name__ == "__main__":
    uvicorn.run(
        app, host="127.0.0.1", port=8000,
        reload=True, log_level="info"
    )